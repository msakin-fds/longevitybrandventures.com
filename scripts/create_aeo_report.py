#!/usr/bin/env python3
"""AEO Audit Report Generator - Creates professional .docx report"""

import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from wp_manager import WordPressManager

class AEOReportGenerator:
    def __init__(self):
        self.doc = Document()
        self.wp = WordPressManager()
        self.setup_styles()

    def setup_styles(self):
        """Configure document styles"""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def add_heading(self, text, level=1, color=(31, 78, 121)):
        """Add formatted heading"""
        heading = self.doc.add_heading(text, level=level)
        heading_format = heading.runs[0]
        heading_format.font.color.rgb = RGBColor(*color)
        heading_format.font.bold = True
        return heading

    def add_key_value(self, key, value, color_key=True):
        """Add key-value pair"""
        p = self.doc.add_paragraph()
        key_run = p.add_run(f"{key}: ")
        key_run.bold = True
        if color_key:
            key_run.font.color.rgb = RGBColor(31, 78, 121)
        p.add_run(str(value))

    def add_recommendation(self, title, description, impact, fix_type, severity='Medium'):
        """Add a recommendation with frontend/backend indicator"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # Title with severity
        title_run = p.add_run(f"• {title}")
        title_run.bold = True
        title_run.font.size = Pt(11)

        severity_colors = {'Critical': (192, 0, 0), 'High': (255, 0, 0), 'Medium': (255, 165, 0), 'Low': (0, 128, 0)}
        severity_run = p.add_run(f"  [{severity}]")
        severity_run.font.color.rgb = RGBColor(*severity_colors.get(severity, (128, 128, 128)))
        severity_run.bold = True

        # Description
        self.doc.add_paragraph(description, style='List Bullet')

        # Impact
        impact_p = self.doc.add_paragraph()
        impact_run = impact_p.add_run("Impact: ")
        impact_run.bold = True
        impact_run.font.color.rgb = RGBColor(31, 78, 121)
        impact_p.add_run(impact)

        # Fix Type
        fix_p = self.doc.add_paragraph()
        fix_run = fix_p.add_run(f"Implementation: ")
        fix_run.bold = True

        if fix_type == 'Frontend':
            type_run = fix_p.add_run(f"{fix_type} ")
            type_run.font.color.rgb = RGBColor(0, 102, 204)
            type_run.bold = True
            fix_p.add_run("(Client-visible changes, affects how visitors see the site)")
        elif fix_type == 'Backend':
            type_run = fix_p.add_run(f"{fix_type} ")
            type_run.font.color.rgb = RGBColor(102, 51, 153)
            type_run.bold = True
            fix_p.add_run("(Server-side changes, improves performance and AI visibility)")
        else:
            type_run = fix_p.add_run(f"{fix_type} ")
            type_run.font.color.rgb = RGBColor(0, 153, 76)
            type_run.bold = True
            fix_p.add_run("(Both frontend and backend improvements needed)")

    def generate_report(self):
        """Generate the complete AEO audit report"""

        # Cover page
        title = self.doc.add_paragraph()
        title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.paragraph_format.space_after = Pt(24)
        title_run = title.add_run("AEO AUDIT REPORT")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 121)

        subtitle = self.doc.add_paragraph()
        subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle.paragraph_format.space_after = Pt(6)
        subtitle_run = subtitle.add_run("AI Engine Optimization Analysis")
        subtitle_run.font.size = Pt(16)

        client = self.doc.add_paragraph()
        client.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        client.paragraph_format.space_after = Pt(24)
        client_run = client.add_run("Longevity Brand Ventures LLC")
        client_run.font.size = Pt(14)
        client_run.font.bold = True

        # Company info
        agency = self.doc.add_paragraph()
        agency.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        agency.paragraph_format.space_after = Pt(6)
        agency_run = agency.add_run("Prepared by: Fresh Design Studio")
        agency_run.font.size = Pt(11)
        agency_run.italic = True

        date = self.doc.add_paragraph()
        date.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date.paragraph_format.space_after = Pt(36)
        date_run = date.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(11)

        self.doc.add_page_break()

        # Executive Summary
        self.add_heading("Executive Summary", level=1)
        self.doc.add_paragraph(
            "This report analyzes your website's optimization for AI Engines (AEO) - ensuring your content is discoverable and featured in AI-generated answers from ChatGPT, Claude, Gemini, Perplexity, and other AI tools. With increasing traffic from AI applications, AEO is now as important as SEO."
        )

        self.add_heading("What is AEO?", level=2)
        self.doc.add_paragraph(
            "AEO (AI Engine Optimization) ensures your website appears in AI-generated responses and summaries. This includes:"
        )
        self.doc.add_paragraph("AI Overviews (Google's AI-powered search summaries)", style='List Bullet')
        self.doc.add_paragraph("AI citations in ChatGPT, Claude, Perplexity, and other LLMs", style='List Bullet')
        self.doc.add_paragraph("Direct answers from AI assistants", style='List Bullet')
        self.doc.add_paragraph("Featured content in AI-powered research tools", style='List Bullet')

        # Site Overview
        self.add_heading("Site Overview", level=1)
        self.add_key_value("Site Name", "Longevity Brand Ventures LLC")
        self.add_key_value("URL", "https://longevitybrandventures.com/")
        self.add_key_value("WordPress Version", "6.9.4")
        self.add_key_value("Theme", "Phlox Pro Child")
        self.add_key_value("Active Plugins", "16")
        self.add_key_value("Inactive Plugins", "2")

        self.doc.add_paragraph()

        # Key Findings
        self.add_heading("Key Findings & Recommendations", level=1)

        self.add_heading("1. SEO Plugin Status", level=2)
        self.add_recommendation(
            "Activate Yoast SEO Plugin",
            "Yoast SEO is currently inactive but installed. This is one of the most important plugins for AEO as it helps optimize content structure, metadata, and readability for AI systems.",
            "Significantly improves AI visibility. Yoast helps structure content in ways AI tools prefer to cite.",
            "Backend",
            "Critical"
        )

        self.add_heading("2. Content Structure & Schema", level=2)
        self.add_recommendation(
            "Implement Rich Schema Markup",
            "Add structured data (Schema.org markup) to your pages. AI tools use this data to understand and cite content accurately. Include: Organization schema, Article schema for blog posts, BreadcrumbList for navigation.",
            "Makes your content more identifiable to AI systems. Increases likelihood of being cited in AI responses.",
            "Backend",
            "Critical"
        )

        self.add_heading("3. Content Quality for AI", level=2)
        self.add_recommendation(
            "Optimize Content for AI Readability",
            "Structure content with clear headings (H1, H2, H3), short paragraphs, bullet points, and definitions. AI tools scan for this structure when deciding what to cite.",
            "Improves chances of content being selected for AI-generated summaries. Better structure = more citations.",
            "Frontend",
            "High"
        )

        self.add_recommendation(
            "Add Authoritative Author Information",
            "Include author bios and credentials on posts. AI systems value author expertise when deciding what to cite.",
            "Establishes authority. AI tools prefer citing content from verified experts.",
            "Frontend",
            "High"
        )

        self.add_heading("4. Meta Information Optimization", level=2)
        self.add_recommendation(
            "Optimize Meta Descriptions",
            "Ensure all pages have unique, descriptive meta descriptions (150-160 characters). These are used by AI tools to understand page content at a glance.",
            "Helps AI systems understand page purpose. Affects how your content is summarized in AI responses.",
            "Backend",
            "Medium"
        )

        self.add_recommendation(
            "Set Canonical URLs",
            "Use canonical tags to prevent duplicate content issues. AI tools may avoid citing duplicate content.",
            "Ensures AI tools cite your preferred version of content, not duplicates.",
            "Backend",
            "Medium"
        )

        self.add_heading("5. Site Performance & Crawlability", level=2)
        self.add_recommendation(
            "Optimize Page Load Speed",
            "AI crawlers scan faster when pages load quickly. Current setup has SG CachePress active (good), but ensure images are optimized, lazy loading is enabled, and minification is configured.",
            "Faster pages get crawled more frequently by AI bots. Impacts indexing in AI training data.",
            "Backend",
            "High"
        )

        self.add_recommendation(
            "Ensure Robots.txt & Sitemap Optimization",
            "Verify your robots.txt allows AI crawlers (GPTBot, Anthropic, Googlebot, etc.). Ensure XML sitemap is present and updated regularly.",
            "Without proper crawlability, AI tools cannot discover your content for citation.",
            "Backend",
            "Critical"
        )

        self.add_heading("6. Mobile Optimization", level=2)
        self.add_recommendation(
            "Verify Mobile Responsiveness",
            "AI tools increasingly use mobile-first indexing. Ensure all content is accessible and readable on mobile devices.",
            "AI systems prioritize mobile-friendly content. Poor mobile experience can hurt discoverability.",
            "Frontend",
            "High"
        )

        self.add_heading("7. Plugin Optimization", level=2)
        self.add_recommendation(
            "Update Elementor to Latest Version",
            "Elementor has an update available (3.35.7 → 4.0.3). Latest versions improve schema markup generation.",
            "Newer versions better support AI-friendly markup generation.",
            "Backend",
            "Medium"
        )

        self.add_recommendation(
            "Update SG Security Plugin",
            "SG Security has an update (1.5.9 → 1.6.0). Keeps security measures current without impacting AI crawlers.",
            "Ensures security doesn't block legitimate AI crawlers (OpenAI, Anthropic, Google, etc.).",
            "Backend",
            "Medium"
        )

        self.add_recommendation(
            "Deactivate Unused Plugins",
            "Master Slider, WordPress Starter, and WP Ulike are inactive. Consider removing them to reduce overhead and improve security.",
            "Unused plugins increase attack surface and slow down site performance.",
            "Backend",
            "Low"
        )

        self.add_heading("8. E-E-A-T Signals (Google Requirement)", level=2)
        self.add_recommendation(
            "Strengthen Author Credentials",
            "Google and AI tools look for Experience, Expertise, Authority, and Trustworthiness. Ensure content reflects these signals.",
            "Critical for health/wellness content. Stronger E-E-A-T = higher visibility in AI summaries.",
            "Frontend",
            "Critical"
        )

        self.add_recommendation(
            "Build External Links & Citations",
            "Get cited by reputable sources in the longevity/health space. Link to authoritative sources from your content.",
            "AI tools consider link patterns when deciding what to cite. More credible sources = more AI citations.",
            "Frontend",
            "High"
        )

        # Implementation Roadmap
        self.add_heading("Implementation Roadmap", level=1)

        self.add_heading("Phase 1: Critical (Week 1-2)", level=2)
        self.doc.add_paragraph("Activate Yoast SEO plugin", style='List Number')
        self.doc.add_paragraph("Verify robots.txt allows all AI crawlers (GPTBot, CCBot, Anthropic, etc.)", style='List Number')
        self.doc.add_paragraph("Add Schema.org markup (Organization, Article)", style='List Number')
        self.doc.add_paragraph("Strengthen author credentials on all content", style='List Number')

        self.add_heading("Phase 2: High Priority (Week 2-3)", level=2)
        self.doc.add_paragraph("Optimize content structure (headings, paragraphs, bullet points)", style='List Number')
        self.doc.add_paragraph("Update Elementor to v4.0.3", style='List Number')
        self.doc.add_paragraph("Optimize all meta descriptions", style='List Number')
        self.doc.add_paragraph("Verify mobile responsiveness across all pages", style='List Number')
        self.doc.add_paragraph("Optimize images and enable lazy loading", style='List Number')

        self.add_heading("Phase 3: Medium Priority (Week 3-4)", level=2)
        self.doc.add_paragraph("Update SG Security plugin", style='List Number')
        self.doc.add_paragraph("Add canonical URLs where needed", style='List Number')
        self.doc.add_paragraph("Set up monitoring for AI crawler activity", style='List Number')

        self.add_heading("Phase 4: Ongoing", level=2)
        self.doc.add_paragraph("Monitor site in AI tools (ChatGPT, Claude, Perplexity, Google AI Overviews)", style='List Number')
        self.doc.add_paragraph("Build high-quality backlinks from authoritative health/wellness sources", style='List Number')
        self.doc.add_paragraph("Keep WordPress and plugins updated", style='List Number')
        self.doc.add_paragraph("Monitor for new AEO best practices", style='List Number')

        # Color Legend
        self.add_heading("Understanding This Report", level=1)

        legend_p = self.doc.add_paragraph()
        legend_p.add_run("Implementation Type:\n").bold = True

        impl_p = self.doc.add_paragraph()
        impl_run = impl_p.add_run("Frontend  ")
        impl_run.font.color.rgb = RGBColor(0, 102, 204)
        impl_run.bold = True
        impl_p.add_run("Changes visible to visitors. Affects how your site looks and feels. Improves user experience alongside AEO.")

        impl_p2 = self.doc.add_paragraph()
        impl_run2 = impl_p2.add_run("Backend  ")
        impl_run2.font.color.rgb = RGBColor(102, 51, 153)
        impl_run2.bold = True
        impl_p2.add_run("Server-side changes. Not visible to visitors. Improves AI discoverability, performance, and security.")

        impl_p3 = self.doc.add_paragraph()
        impl_run3 = impl_p3.add_run("Both  ")
        impl_run3.font.color.rgb = RGBColor(0, 153, 76)
        impl_run3.bold = True
        impl_p3.add_run("Improvements affecting both visitors and AI systems.")

        self.doc.add_paragraph()

        severity_p = self.doc.add_paragraph()
        severity_p.add_run("Severity Levels:\n").bold = True

        for severity, desc in [
            ("Critical", "Implement immediately. Significantly impacts AI visibility."),
            ("High", "Implement soon. Notable impact on AEO performance."),
            ("Medium", "Implement after high-priority items. Incremental improvements."),
            ("Low", "Nice to have. Polish and optimization.")
        ]:
            sev_p = self.doc.add_paragraph()
            sev_run = sev_p.add_run(f"{severity}  ")
            sev_run.bold = True
            sev_p.add_run(desc)

        # Footer
        self.doc.add_page_break()
        self.add_heading("Next Steps", level=1)
        self.doc.add_paragraph(
            "Review this report with Fresh Design Studio to prioritize improvements based on your business goals. We recommend starting with Phase 1 items to maximize impact on AI visibility."
        )
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "AEO is an evolving field. As AI tools become more sophisticated, we'll monitor best practices and adjust recommendations accordingly."
        )

        footer = self.doc.add_paragraph()
        footer.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer.paragraph_format.space_before = Pt(24)
        footer_run = footer.add_run("Fresh Design Studio")
        footer_run.italic = True
        footer_run.font.size = Pt(10)

    def save(self, filename='AEO_Audit_Report.docx'):
        """Save the report"""
        self.doc.save(filename)
        return filename

if __name__ == '__main__':
    print('Generating AEO Audit Report...')
    generator = AEOReportGenerator()
    generator.generate_report()
    filename = generator.save()
    print(f'✓ Report saved: {filename}')
