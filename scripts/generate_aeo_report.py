#!/usr/bin/env python3
"""AEO Audit Report Generator for longevitybrandventures.com"""

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1F, 0x4E, 0x79)
DARK_BLUE   = RGBColor(0x1B, 0x63, 0x9A)
LIGHT_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
DARK_GREY   = RGBColor(0x40, 0x40, 0x40)
MID_GREY    = RGBColor(0x70, 0x70, 0x70)
RED         = RGBColor(0xC0, 0x00, 0x00)
ORANGE      = RGBColor(0xE0, 0x70, 0x00)
AMBER       = RGBColor(0xFF, 0xC0, 0x00)
GREEN       = RGBColor(0x37, 0x86, 0x40)
FRONTEND_C  = RGBColor(0x00, 0x70, 0xC0)
BACKEND_C   = RGBColor(0x70, 0x30, 0xA0)
BOTH_C      = RGBColor(0x37, 0x86, 0x40)

SEVERITY_COLOUR = {
    "Critical":  RED,
    "High":      ORANGE,
    "Medium":    AMBER,
    "Low":       GREEN,
}

SEVERITY_BG = {
    "Critical":  "FFE0E0",
    "High":      "FFF0DC",
    "Medium":    "FFFACC",
    "Low":       "E8F5E9",
}

# ── REAL AUDIT DATA ───────────────────────────────────────────────────────────
AUDIT = {
    "site_name":        "Longevity Brand Ventures LLC",
    "site_url":         "https://longevitybrandventures.com",
    "wp_version":       "6.9.4",
    "theme":            "Phlox Pro Child",
    "audit_date":       datetime.now().strftime("%B %d, %Y"),
    "prepared_by":      "Fresh Design Studio",
    "total_plugins":    17,
    "active_plugins":   12,
    "inactive_plugins": 4,
    "page_count":       35,
    "post_count":       12,
    "score":            38,          # out of 100
}

FINDINGS = [
    # ── CRITICAL ──────────────────────────────────────────────────────────────
    {
        "id":       "F01",
        "area":     "Schema Markup",
        "title":    "Zero Structured Data on the Website",
        "severity": "Critical",
        "type":     "Backend",
        "problem":  (
            "The entire site contains no JSON-LD or any Schema.org markup whatsoever. "
            "Structured data is the primary way AI engines (Google AI Overviews, ChatGPT, "
            "Perplexity, Claude, Gemini) understand who you are, what you do, and whether "
            "to cite you in their responses. Without it, the site is nearly invisible to AI."
        ),
        "fix":      (
            "Add Organisation, WebSite, BreadcrumbList, and Article schema blocks "
            "to all key pages. The homepage needs an Organisation block with name, "
            "description, URL, logo, sameAs (social profiles), and foundingDate at minimum."
        ),
        "impact":   "Unlocking AI citations; this is the single biggest AEO gain available.",
        "effort":   "2-4 hours",
    },
    {
        "id":       "F02",
        "area":     "Yoast SEO Plugin",
        "title":    "Yoast SEO Is Installed But Switched Off",
        "severity": "Critical",
        "type":     "Backend",
        "problem":  (
            "Yoast SEO (version 27.1.1, latest 27.4) is installed on the server but is "
            "deactivated. This means the site has no SEO title templates, no meta-description "
            "control, no canonical management, and no XML sitemap generation — all essential "
            "for AEO."
        ),
        "fix":      (
            "Activate Yoast SEO, run the configuration wizard, then update it to v27.4. "
            "Once active, complete all page/post SEO titles and focus keyphrases in Yoast."
        ),
        "impact":   "Enables structured meta data, sitemaps, and content analysis across every page.",
        "effort":   "30 minutes to activate + 2-3 hours to configure all pages",
    },
    {
        "id":       "F03",
        "area":     "Robots.txt",
        "title":    "No robots.txt File Exists",
        "severity": "Critical",
        "type":     "Backend",
        "problem":  (
            "The site has no robots.txt file. This means AI crawler bots (GPTBot for ChatGPT, "
            "Claude-Web, Anthropic-AI, PerplexityBot, Google-Extended for Gemini) receive no "
            "guidance. Some may refuse to index a site with missing robots.txt; others may "
            "over-crawl and trigger server load issues."
        ),
        "fix":      (
            "Create /robots.txt explicitly allowing all major AI crawlers: Googlebot, "
            "GPTBot, ClaudeBot, PerplexityBot, Google-Extended, CCBot, and Amazonbot. "
            "Point to the XML sitemap URL at the bottom of the file."
        ),
        "impact":   "Signals to AI systems that the site is open to being indexed and cited.",
        "effort":   "30 minutes",
    },
    {
        "id":       "F04",
        "area":     "XML Sitemap",
        "title":    "No XML Sitemap Found",
        "severity": "Critical",
        "type":     "Backend",
        "problem":  (
            "There is no sitemap.xml file on the server. AI crawlers and search engines use "
            "sitemaps to discover and prioritise which pages to read. Without one, many of "
            "the 35 pages on this site — including portfolio brand pages — may never be crawled."
        ),
        "fix":      (
            "Activating Yoast SEO (F02) will auto-generate a sitemap at "
            "longevitybrandventures.com/sitemap_index.xml. Submit it to Google Search Console "
            "and Bing Webmaster Tools. Confirm all portfolio pages appear in it."
        ),
        "impact":   "Ensures all 35 pages are discovered by AI crawlers and search engines.",
        "effort":   "15 minutes (auto-generated once Yoast is active)",
    },
    {
        "id":       "F05",
        "area":     "H1 Heading",
        "title":    "Homepage Has No H1 Tag",
        "severity": "Critical",
        "type":     "Frontend",
        "problem":  (
            "A live crawl of the homepage returned zero H1 tags. AI models treat the H1 "
            "as the primary topic signal for a page. Missing H1 means AI tools cannot "
            "confidently determine what the page is about and are unlikely to cite it."
        ),
        "fix":      (
            "Add a clear, keyword-rich H1 to the homepage via Elementor. Example: "
            "\"Connecting Longevity Brand Entrepreneurs with Investors & Advisors\" — "
            "visible to users and crawlers. Only one H1 per page."
        ),
        "impact":   "Immediately tells AI what the site is about; foundational for citations.",
        "effort":   "15 minutes",
    },
    # ── HIGH ──────────────────────────────────────────────────────────────────
    {
        "id":       "F06",
        "area":     "Open Graph & Social Metadata",
        "title":    "No Open Graph or Twitter/X Card Tags",
        "severity": "High",
        "type":     "Backend",
        "problem":  (
            "The site has no og:title, og:description, og:image, og:type, "
            "twitter:card, or twitter:image tags. These tags are read by AI tools "
            "and social platforms when summarising or sharing your content. Their "
            "absence means AI systems receive less context about each page."
        ),
        "fix":      (
            "Yoast SEO (F02) will generate these automatically once configured. "
            "Set a default social sharing image (1200x630px) and customise "
            "OG titles/descriptions for the homepage, About, and all portfolio pages."
        ),
        "impact":   "Improves how the site appears when shared and how AI platforms summarise it.",
        "effort":   "1-2 hours after Yoast is active",
    },
    {
        "id":       "F07",
        "area":     "Meta Description Quality",
        "title":    "Homepage Meta Description Is Unoptimised (450+ Characters)",
        "severity": "High",
        "type":     "Backend",
        "problem":  (
            "The current homepage meta description is 450+ characters long and appears to be "
            "a raw content dump from the page body, beginning: \"Collaboration Between "
            "Entrepreneurs, Investors & Advisors Your Pathway to Connection…\" — this is "
            "truncated in search results and difficult for AI to parse into a clean summary."
        ),
        "fix":      (
            "Write a concise, compelling meta description of 150-160 characters that clearly "
            "explains who Longevity Brand Ventures is and who they serve. Example: "
            "\"Longevity Brand Ventures connects health & wellness entrepreneurs with strategic "
            "investors and advisors in the consumer packaged goods space.\""
        ),
        "impact":   "AI tools use meta descriptions as page summaries; a clean one increases citation accuracy.",
        "effort":   "30 minutes per page",
    },
    {
        "id":       "F08",
        "area":     "Site Tagline",
        "title":    "Site Tagline \"An advantaged organization\" Is Too Vague",
        "severity": "High",
        "type":     "Frontend",
        "problem":  (
            "The WordPress site tagline (visible in the browser tab and used by AI tools "
            "when summarising the site) reads \"An advantaged organization\" — this is generic, "
            "non-descriptive, and provides no keyword signal for AI systems trying to classify "
            "the business."
        ),
        "fix":      (
            "Update in WordPress Admin > Settings > General > Tagline. Suggested: "
            "\"Bridging Longevity Brands with Investors & Strategic Advisors\" or similar "
            "that clearly states the niche and value proposition."
        ),
        "impact":   "AI tools include the site title + tagline when generating citations; a better tagline improves accuracy.",
        "effort":   "5 minutes",
    },
    {
        "id":       "F09",
        "area":     "Blog Content",
        "title":    "Blog Contains 12 Placeholder / Lorem Ipsum Posts from 2019",
        "severity": "High",
        "type":     "Frontend",
        "problem":  (
            "The blog section has 12 published posts dating from October 2019, with titles "
            "such as \"When darkness overspreads my eyes\" and \"Foliage of my trees\" — these "
            "appear to be demo/placeholder content. Published dummy content actively harms AEO: "
            "AI systems index and may cite it, damaging credibility and topical authority."
        ),
        "fix":      (
            "Delete or unpublish all placeholder posts. Replace with at minimum 3-5 "
            "substantive articles about longevity investing, consumer brands, or the firm's "
            "investment thesis — written by a named author with credentials."
        ),
        "impact":   "Removes credibility risk; allows AI to build a coherent topical profile for the site.",
        "effort":   "1 hour to remove; ongoing content creation",
    },
    {
        "id":       "F10",
        "area":     "Elementor Plugin",
        "title":    "Elementor Is One Major Version Behind (3.35.7 vs 4.0.3)",
        "severity": "High",
        "type":     "Backend",
        "problem":  (
            "Elementor is currently on version 3.35.7 while version 4.0.3 is available. "
            "This is a major version jump with performance improvements, Core Web Vitals "
            "optimisations, and better semantic HTML output — all of which affect how AI "
            "crawlers read page content. Running outdated page builders can also introduce "
            "security vulnerabilities."
        ),
        "fix":      "Update Elementor to v4.0.3 via WordPress Admin > Plugins. Back up the site first.",
        "impact":   "Better HTML output, improved Core Web Vitals scores, and reduced security risk.",
        "effort":   "30 minutes (including backup)",
    },
    # ── MEDIUM ────────────────────────────────────────────────────────────────
    {
        "id":       "F11",
        "area":     "E-E-A-T Signals",
        "title":    "No Author Credentials or Expertise Signals on Content",
        "severity": "Medium",
        "type":     "Frontend",
        "problem":  (
            "Google and AI systems evaluate content through E-E-A-T: Experience, Expertise, "
            "Authoritativeness, and Trustworthiness. This is especially critical in finance "
            "and health/longevity niches. The site has no visible author bios, no credentials "
            "listed, no team page showcasing expertise, and no LinkedIn or professional profile links."
        ),
        "fix":      (
            "Add a Team or About page with photos, titles, and credentials for key people. "
            "Link to LinkedIn profiles. Add an author bio block to all blog posts and "
            "portfolio pages. Include the firm's founding year and track record where possible."
        ),
        "impact":   "Directly affects whether AI tools treat the site as a credible, citable source.",
        "effort":   "4-6 hours",
    },
    {
        "id":       "F12",
        "area":     "Content Architecture",
        "title":    "Multiple Duplicate & Outdated Draft Pages Pollute Site Structure",
        "severity": "Medium",
        "type":     "Backend",
        "problem":  (
            "The site contains several legacy and duplicate pages: \"Portfolio Old\", "
            "\"Portfolio - NEW\" (draft), \"Blog Old\", \"Greg Old\", \"Home\" (draft), "
            "\"Contact Old\", \"About Old\", and \"Sample Page\" — 8 junk pages that "
            "dilute crawl budget and confuse AI about the site's canonical content."
        ),
        "fix":      (
            "Delete all pages with 'Old', 'Old', and 'Sample Page' in the title. "
            "Ensure draft pages are not indexable. Set a clear page hierarchy with "
            "Home, About, Portfolio, Investment Thesis, and Contact as the five core pages."
        ),
        "impact":   "Focuses crawl budget on meaningful pages; reduces AI confusion about site structure.",
        "effort":   "1 hour",
    },
    {
        "id":       "F13",
        "area":     "Security Plugin",
        "title":    "SG Security Plugin Needs Updating (1.5.9 to 1.6.0)",
        "severity": "Medium",
        "type":     "Backend",
        "problem":  (
            "The SiteGround Security plugin is one version behind. Keeping security plugins "
            "current prevents vulnerabilities that could lead to site blacklisting — which "
            "would cause AI tools to stop citing the site entirely."
        ),
        "fix":      "Update SG Security via WordPress Admin > Plugins > Update.",
        "impact":   "Reduces risk of site being flagged or blacklisted by security databases.",
        "effort":   "5 minutes",
    },
    {
        "id":       "F14",
        "area":     "Administrator Accounts",
        "title":    "Five Administrator Accounts — Excess Admin Roles Are a Security Risk",
        "severity": "Medium",
        "type":     "Backend",
        "problem":  (
            "The site has 5 WordPress accounts with full Administrator access "
            "(anson.wu@freshds.com, fdsthinker@fds.com, kendra@voxagency.co, "
            "msakin@freshds.com, Webmaster). Multiple admin accounts widen the attack surface. "
            "A compromised account could inject harmful content or malware, causing AI tools "
            "to stop citing the site."
        ),
        "fix":      (
            "Review all admin accounts. Downgrade any account that does not need full admin "
            "access to Editor or Author role. Remove the generic \"Webmaster\" account if unused. "
            "Enable two-factor authentication (2FA) for all remaining admins."
        ),
        "impact":   "Reduces security exposure; protects site integrity and AI citation standing.",
        "effort":   "1 hour",
    },
    {
        "id":       "F15",
        "area":     "Inactive Plugins",
        "title":    "4 Inactive Plugins Left Installed on the Server",
        "severity": "Medium",
        "type":     "Backend",
        "problem":  (
            "Master Slider, WordPress Starter, WP Ulike, and Mailchimp for WP are installed "
            "but inactive. Inactive plugins still represent attack vectors if their files "
            "contain vulnerabilities. Mailchimp for WP also has an available update (4.12.2) "
            "suggesting it hasn't been maintained."
        ),
        "fix":      (
            "Delete Master Slider, WordPress Starter, and WP Ulike if not needed. "
            "If Mailchimp integration is needed, update and reactivate it; otherwise delete it."
        ),
        "impact":   "Reduces security surface area and keeps the codebase lean.",
        "effort":   "30 minutes",
    },
    # ── LOW ───────────────────────────────────────────────────────────────────
    {
        "id":       "F16",
        "area":     "Image Optimisation",
        "title":    "No Image Optimisation Plugin Active",
        "severity": "Low",
        "type":     "Backend",
        "problem":  (
            "There is no active image compression or WebP conversion plugin. "
            "Slow-loading images hurt Core Web Vitals (LCP score), which AI crawlers "
            "like Google's use as a quality signal when ranking content for AI Overviews."
        ),
        "fix":      (
            "Install and configure a free plugin such as Smush or ShortPixel to automatically "
            "compress images and serve WebP format. SG CachePress (already active) has some "
            "image optimisation settings — ensure these are fully enabled."
        ),
        "impact":   "Improves page speed scores, which is a ranking factor for AI Overviews.",
        "effort":   "1-2 hours",
    },
    {
        "id":       "F17",
        "area":     "Internal Linking",
        "title":    "No Clear Internal Linking Strategy Visible",
        "severity": "Low",
        "type":     "Frontend",
        "problem":  (
            "AI tools follow internal links to understand site hierarchy and content "
            "relationships. Without a deliberate internal linking structure connecting "
            "portfolio brand pages back to the Investment Thesis, About, and Blog pages, "
            "the site appears as a flat collection of disconnected pages."
        ),
        "fix":      (
            "Add contextual links on each portfolio page pointing to the Investment Thesis "
            "and About pages. Add links from the Blog to related portfolio brands. "
            "Ensure the main navigation mirrors the priority of pages for AI crawlers."
        ),
        "impact":   "Helps AI understand the relationship between content; reinforces topical authority.",
        "effort":   "2-3 hours",
    },
]

ROADMAP = [
    {
        "phase": "Phase 1 — Immediate (Week 1)",
        "colour": "C00000",
        "items": [
            ("F01", "Add Organisation + WebSite Schema to homepage", "Backend"),
            ("F02", "Activate and update Yoast SEO plugin", "Backend"),
            ("F03", "Create robots.txt allowing all AI crawlers", "Backend"),
            ("F04", "Confirm sitemap is live and submit to Google Search Console", "Backend"),
            ("F05", "Add H1 heading to homepage via Elementor", "Frontend"),
        ],
    },
    {
        "phase": "Phase 2 — High Priority (Week 2)",
        "colour": "E07000",
        "items": [
            ("F06", "Configure Open Graph and Twitter card tags via Yoast", "Backend"),
            ("F07", "Rewrite all page meta descriptions (150-160 chars)", "Backend"),
            ("F08", "Update WordPress site tagline to reflect the business niche", "Frontend"),
            ("F09", "Delete placeholder blog posts; plan real content", "Frontend"),
            ("F10", "Update Elementor to v4.0.3 (backup first)", "Backend"),
        ],
    },
    {
        "phase": "Phase 3 — Medium Priority (Week 3)",
        "colour": "375E97",
        "items": [
            ("F11", "Add author bios and team credentials to all content", "Frontend"),
            ("F12", "Delete legacy/duplicate pages; clean up page structure", "Backend"),
            ("F13", "Update SG Security plugin to v1.6.0", "Backend"),
            ("F14", "Review admin accounts; enable 2FA for all administrators", "Backend"),
            ("F15", "Remove unused inactive plugins", "Backend"),
        ],
    },
    {
        "phase": "Phase 4 — Ongoing",
        "colour": "378640",
        "items": [
            ("F16", "Install and configure image optimisation plugin", "Backend"),
            ("F17", "Build internal linking between portfolio, thesis, and blog pages", "Frontend"),
            ("-",   "Publish 1-2 expert articles per month on longevity investing", "Frontend"),
            ("-",   "Monitor AI mentions in ChatGPT, Perplexity, Google AI Overviews", "Both"),
            ("-",   "Review and refresh schema markup quarterly", "Backend"),
        ],
    },
]


# ── HELPERS ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_run(para, text, bold=False, italic=False, size=None, colour=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if size:
        run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    return run

def add_para(doc, text='', alignment=None, space_before=0, space_after=6):
    p = doc.add_paragraph(text)
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p

def add_hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'),  '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pb.append(bottom)
    pPr.append(pb)

def set_page_margins(doc, top=1.5, bottom=1.5, left=2.0, right=2.0):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


# ── BUILDER ──────────────────────────────────────────────────────────────────

def build_report():
    doc = Document()
    set_page_margins(doc)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # Top accent bar (paragraph with shading)
    top_bar = doc.add_paragraph()
    top_bar.paragraph_format.space_before = Pt(0)
    top_bar.paragraph_format.space_after  = Pt(0)
    top_bar_run = top_bar.add_run("   ")
    top_bar_run.font.size = Pt(4)
    pPr = top_bar._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F4E79')
    pPr.append(shd)

    doc.add_paragraph().paragraph_format.space_after = Pt(48)

    # Report type label
    lbl = add_para(doc, space_before=0, space_after=6)
    lbl.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(lbl, "AI ENGINE OPTIMISATION AUDIT", bold=True, size=11, colour=LIGHT_BLUE)

    # Main title
    title_p = add_para(doc, space_before=0, space_after=8)
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(title_p, "AEO AUDIT REPORT", bold=True, size=32, colour=NAVY)

    # Client name
    client_p = add_para(doc, space_before=4, space_after=4)
    client_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(client_p, AUDIT["site_name"], bold=True, size=18, colour=DARK_BLUE)

    # URL
    url_p = add_para(doc, space_before=0, space_after=48)
    url_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(url_p, AUDIT["site_url"], italic=True, size=11, colour=MID_GREY)

    # Overall score box
    score_p = add_para(doc, space_before=0, space_after=4)
    score_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(score_p, "Overall AEO Score", bold=True, size=13, colour=DARK_GREY)

    score_val = add_para(doc, space_before=0, space_after=6)
    score_val.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(score_val, f"{AUDIT['score']}/100", bold=True, size=42, colour=RED)

    score_note = add_para(doc, space_before=0, space_after=48)
    score_note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(score_note, "Significant improvements required to achieve strong AI visibility.",
            italic=True, size=10, colour=MID_GREY)

    # Meta info
    meta_p = add_para(doc, space_before=0, space_after=4)
    meta_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(meta_p, f"Prepared by:  ", bold=True, size=11, colour=DARK_GREY)
    add_run(meta_p, AUDIT["prepared_by"], size=11, colour=DARK_GREY)

    date_p = add_para(doc, space_before=0, space_after=4)
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(date_p, f"Date:  ", bold=True, size=11, colour=DARK_GREY)
    add_run(date_p, AUDIT["audit_date"], size=11, colour=DARK_GREY)

    confidential_p = add_para(doc, space_before=24, space_after=0)
    confidential_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(confidential_p, "CONFIDENTIAL", bold=True, size=9, colour=MID_GREY)

    doc.add_page_break()

    # ── WHAT IS AEO ──────────────────────────────────────────────────────────
    h = doc.add_heading("What Is AEO?", level=1)
    h.runs[0].font.color.rgb = NAVY

    intro_text = (
        "Search is changing. Millions of people now get answers directly from AI tools — "
        "Google AI Overviews, ChatGPT, Perplexity, Claude, and Gemini — without ever clicking "
        "a link. AI Engine Optimisation (AEO) ensures that when someone asks an AI about topics "
        "relevant to your business, your website is the source it cites or summarises."
    )
    add_para(doc, intro_text, space_after=8)

    add_para(doc,
        "AEO is distinct from traditional SEO. While SEO targets rankings on a results page, "
        "AEO targets inclusion in AI-generated answers. The two complement each other: a site "
        "strong in SEO gains a head start on AEO, but additional optimisations — schema markup, "
        "content authority signals, and AI crawler access — are required to fully benefit.",
        space_after=8)

    add_hrule(doc)

    # ── SITE SNAPSHOT ────────────────────────────────────────────────────────
    h2 = doc.add_heading("Site Snapshot", level=1)
    h2.runs[0].font.color.rgb = NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    snap_table = doc.add_table(rows=9, cols=2)
    snap_table.style = 'Table Grid'
    snap_rows = [
        ("Website URL",           AUDIT["site_url"]),
        ("Site Name",             AUDIT["site_name"]),
        ("WordPress Version",     AUDIT["wp_version"]),
        ("Active Theme",          AUDIT["theme"]),
        ("Active Plugins",        str(AUDIT["active_plugins"])),
        ("Inactive Plugins",      str(AUDIT["inactive_plugins"])),
        ("Total Published Pages", str(AUDIT["page_count"])),
        ("Blog Posts",            str(AUDIT["post_count"]) + " (all from 2019 — placeholder content)"),
        ("Audit Date",            AUDIT["audit_date"]),
    ]
    for i, (label, value) in enumerate(snap_rows):
        row = snap_table.rows[i]
        set_cell_bg(row.cells[0], "1F4E79")
        p0 = row.cells[0].paragraphs[0]
        add_run(p0, label, bold=True, size=10, colour=WHITE)
        p1 = row.cells[1].paragraphs[0]
        add_run(p1, value, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_hrule(doc)

    # ── SCORING SUMMARY ───────────────────────────────────────────────────────
    h3 = doc.add_heading("Audit Score Breakdown", level=1)
    h3.runs[0].font.color.rgb = NAVY

    score_table = doc.add_table(rows=6, cols=4)
    score_table.style = 'Table Grid'

    headers = ["AEO Category", "Max Score", "Your Score", "Status"]
    for i, hdr in enumerate(headers):
        cell = score_table.rows[0].cells[i]
        set_cell_bg(cell, "2E75B6")
        add_run(cell.paragraphs[0], hdr, bold=True, size=10, colour=WHITE)

    score_data = [
        ("Structured Data / Schema",   "25",  "0",  "Critical — Missing"),
        ("Crawlability & Sitemaps",     "20",  "5",  "Critical — No robots.txt or sitemap"),
        ("Content Quality & Structure", "25",  "13", "Needs Work"),
        ("Technical SEO Foundations",  "20",  "12", "Needs Work"),
        ("Authority & Trust Signals",  "10",  "8",  "Partially Present"),
    ]
    status_colours = {
        "Critical — Missing":                  "FFD7D7",
        "Critical — No robots.txt or sitemap": "FFD7D7",
        "Needs Work":                          "FFF0D0",
        "Partially Present":                   "FFFACC",
    }
    for r_idx, (cat, max_s, your_s, status) in enumerate(score_data, start=1):
        row = score_table.rows[r_idx]
        add_run(row.cells[0].paragraphs[0], cat, size=10)
        add_run(row.cells[1].paragraphs[0], max_s, size=10)
        add_run(row.cells[2].paragraphs[0], your_s, bold=True, size=10,
                colour=RED if int(your_s) < int(max_s) // 2 else ORANGE)
        set_cell_bg(row.cells[3], status_colours.get(status, "FFFFFF"))
        add_run(row.cells[3].paragraphs[0], status, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_hrule(doc)

    # ── LEGEND ───────────────────────────────────────────────────────────────
    leg_head = doc.add_heading("Reading This Report", level=1)
    leg_head.runs[0].font.color.rgb = NAVY

    legend_table = doc.add_table(rows=8, cols=3)
    legend_table.style = 'Table Grid'

    leg_headers = ["Label", "Colour", "Meaning"]
    for i, hdr in enumerate(leg_headers):
        cell = legend_table.rows[0].cells[i]
        set_cell_bg(cell, "2E75B6")
        add_run(cell.paragraphs[0], hdr, bold=True, size=10, colour=WHITE)

    legend_rows = [
        ("Frontend Fix",  "0070C0", "This change is VISIBLE to your website visitors. It affects how the site looks, reads, or feels."),
        ("Backend Fix",   "7030A0", "This change is NOT visible to visitors. It happens behind the scenes — code, settings, or server config."),
        ("Both",          "378640", "Both visible and behind-the-scenes changes are required."),
        ("CRITICAL",      "C00000", "Implement in Week 1. Significantly blocks AI visibility right now."),
        ("HIGH",          "E07000", "Implement in Week 2. Notable negative impact if left unfixed."),
        ("MEDIUM",        "C07800", "Implement in Week 3. Incremental improvements."),
        ("LOW",           "378640", "Implement in Phase 4. Polish and long-term optimisation."),
    ]
    for r_idx, (label, colour, meaning) in enumerate(legend_rows, start=1):
        row = legend_table.rows[r_idx]
        set_cell_bg(row.cells[0], colour)
        add_run(row.cells[0].paragraphs[0], label, bold=True, size=10, colour=WHITE)
        set_cell_bg(row.cells[1], colour + "40"[::-1])   # lighter shade
        add_run(row.cells[2].paragraphs[0], meaning, size=10)

    doc.add_page_break()

    # ── FINDINGS ─────────────────────────────────────────────────────────────
    h_find = doc.add_heading("Detailed Findings & Recommendations", level=1)
    h_find.runs[0].font.color.rgb = NAVY

    type_colours = {
        "Frontend": "0070C0",
        "Backend":  "7030A0",
        "Both":     "378640",
    }

    for f in FINDINGS:
        sev = f["severity"]
        ftype = f["type"]

        # Finding header table (ID | Severity | Type)
        hdr_tbl = doc.add_table(rows=1, cols=3)
        hdr_tbl.style = 'Table Grid'
        hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        c0 = hdr_tbl.rows[0].cells[0]
        set_cell_bg(c0, "1F4E79")
        add_run(c0.paragraphs[0], f["id"], bold=True, size=10, colour=WHITE)

        c1 = hdr_tbl.rows[0].cells[1]
        set_cell_bg(c1, SEVERITY_BG[sev])
        add_run(c1.paragraphs[0], sev.upper(), bold=True, size=10,
                colour=SEVERITY_COLOUR[sev])

        c2 = hdr_tbl.rows[0].cells[2]
        set_cell_bg(c2, type_colours[ftype] + "20")
        add_run(c2.paragraphs[0],
                f"{ftype} Fix",
                bold=True, size=10,
                colour=RGBColor(*bytes.fromhex(type_colours[ftype])))

        # Title
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(2)
        title_p.paragraph_format.space_after  = Pt(4)
        add_run(title_p, f["area"] + "  —  ", bold=True, size=11, colour=DARK_BLUE)
        add_run(title_p, f["title"], bold=True, size=11, colour=BLACK)

        # Problem
        prob_p = doc.add_paragraph()
        prob_p.paragraph_format.space_before = Pt(0)
        prob_p.paragraph_format.space_after  = Pt(3)
        add_run(prob_p, "Issue:  ", bold=True, size=10, colour=RED)
        add_run(prob_p, f["problem"], size=10)

        # Fix
        fix_p = doc.add_paragraph()
        fix_p.paragraph_format.space_before = Pt(0)
        fix_p.paragraph_format.space_after  = Pt(3)
        add_run(fix_p, "Recommended Fix:  ", bold=True, size=10, colour=GREEN)
        add_run(fix_p, f["fix"], size=10)

        # Impact + Effort inline
        ie_p = doc.add_paragraph()
        ie_p.paragraph_format.space_before = Pt(0)
        ie_p.paragraph_format.space_after  = Pt(10)
        add_run(ie_p, "Expected Impact:  ", bold=True, size=10, colour=DARK_BLUE)
        add_run(ie_p, f["impact"], size=10)
        add_run(ie_p, "     Effort:  ", bold=True, size=10, colour=DARK_GREY)
        add_run(ie_p, f["effort"], size=10)

    add_hrule(doc)

    # ── ROADMAP ───────────────────────────────────────────────────────────────
    h_road = doc.add_heading("Implementation Roadmap", level=1)
    h_road.runs[0].font.color.rgb = NAVY

    add_para(doc,
        "The following roadmap prioritises findings by impact and effort. Items in Phase 1 "
        "should be completed before moving to Phase 2. Each phase builds on the last.",
        space_after=10)

    for phase in ROADMAP:
        # Phase header
        ph_p = doc.add_paragraph()
        ph_p.paragraph_format.space_before = Pt(8)
        ph_p.paragraph_format.space_after  = Pt(4)
        add_run(ph_p, phase["phase"], bold=True, size=12,
                colour=RGBColor(*bytes.fromhex(phase["colour"])))

        # Items table
        road_tbl = doc.add_table(rows=len(phase["items"]) + 1, cols=4)
        road_tbl.style = 'Table Grid'

        for i, col_name in enumerate(["Ref", "Action", "Type", "Done"]):
            c = road_tbl.rows[0].cells[i]
            set_cell_bg(c, phase["colour"])
            add_run(c.paragraphs[0], col_name, bold=True, size=9, colour=WHITE)

        for r, (ref, action, itype) in enumerate(phase["items"], start=1):
            row = road_tbl.rows[r]
            add_run(row.cells[0].paragraphs[0], ref,    size=9)
            add_run(row.cells[1].paragraphs[0], action, size=9)
            c_type = row.cells[2]
            type_hex = type_colours.get(itype, "555555")
            set_cell_bg(c_type, type_hex + "20")
            add_run(c_type.paragraphs[0], itype, bold=True, size=9,
                    colour=RGBColor(*bytes.fromhex(type_hex)))
            add_run(row.cells[3].paragraphs[0], "[ ]", size=9)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_hrule(doc)

    # ── CLOSING ───────────────────────────────────────────────────────────────
    h_next = doc.add_heading("Next Steps", level=1)
    h_next.runs[0].font.color.rgb = NAVY

    add_para(doc,
        "We recommend scheduling a kickoff call with Fresh Design Studio to review Phase 1 "
        "priorities. Most Phase 1 items can be completed within a single focused session. "
        "Phase 2 and 3 items can then be tackled in subsequent sprints.",
        space_after=8)

    add_para(doc,
        "AEO is a fast-moving discipline. As AI tools evolve, best practices shift. "
        "Fresh Design Studio will monitor changes in how major AI platforms crawl and cite "
        "content, and will update recommendations accordingly.",
        space_after=24)

    footer_p = add_para(doc, space_before=24, space_after=0)
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(footer_p, "Fresh Design Studio", bold=True, size=12, colour=NAVY)

    footer2_p = add_para(doc, space_before=2, space_after=0)
    footer2_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(footer2_p, "Confidential — prepared exclusively for Longevity Brand Ventures LLC",
            italic=True, size=9, colour=MID_GREY)

    return doc


if __name__ == "__main__":
    print("Generating AEO Audit Report...")
    doc = build_report()
    out = "LongevityBrandVentures_AEO_Audit_Report.docx"
    doc.save(out)
    print(f"Report saved: {out}")
